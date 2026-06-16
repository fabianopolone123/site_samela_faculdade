from django.conf import settings
from django.core import mail
from django.test import TestCase
from django.urls import reverse
from docx import Document
from io import BytesIO

from .forms import FieldMigrationForm, TopicFieldForm
from .models import AllowedSignupEmail, AuditLog, CostField, CostRecord, CostRecordValue, CostTopic, SignupCode, User
from .views import build_budget_ready_context, build_fapesp_export_pages


class SignupFlowTests(TestCase):
    def test_login_page_displays_ptbr_texts_correctly(self):
        response = self.client.get(reverse('login'))

        self.assertContains(response, 'Instituição')
        self.assertContains(response, 'Código')
        self.assertNotContains(response, 'InstituiÃ§Ã£o', html=False)
        self.assertNotContains(response, 'CÃ³digo', html=False)

    def test_disallowed_email_cannot_start_signup(self):
        response = self.client.post(
            reverse('signup_email'),
            {'email': 'nao-autorizado@example.com'},
        )

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'não está autorizado', html=False)
        self.assertEqual(SignupCode.objects.count(), 0)

    def test_full_signup_flow_creates_user_and_sends_emails(self):
        email = next(iter(settings.ALLOWED_SIGNUP_EMAILS))

        response = self.client.post(reverse('signup_email'), {'email': email})
        self.assertRedirects(response, reverse('login'))
        self.assertEqual(len(mail.outbox), 1)

        code = SignupCode.objects.get(email=email)
        response = self.client.post(reverse('signup_code'), {'code': code.code})
        self.assertRedirects(response, reverse('login'))

        response = self.client.post(
            reverse('signup_password'),
            {'password1': 'SenhaSegura123', 'password2': 'SenhaSegura123'},
        )
        self.assertRedirects(response, reverse('login'))
        self.assertTrue(User.objects.filter(email=email).exists())
        self.assertEqual(len(mail.outbox), 2)


class SeededAccessTests(TestCase):
    def test_seeded_test_user_can_login_with_login_name(self):
        response = self.client.post(
            reverse('login'),
            {'username': 'fabiano', 'password': '123'},
        )

        self.assertRedirects(response, reverse('dashboard'))


class DynamicTopicBudgetTests(TestCase):
    def setUp(self):
        self.client.post(reverse('login'), {'username': 'fabiano', 'password': '123'})

    def test_default_topics_and_fields_are_seeded(self):
        self.assertTrue(
            CostTopic.objects.filter(
                name='Material permanente adquirido no país e importado'
            ).exists()
        )
        self.assertTrue(
            CostTopic.objects.filter(
                name='Material de consumo adquirido no país e importado'
            ).exists()
        )
        topic = CostTopic.objects.get(name='Material permanente adquirido no país e importado')
        self.assertTrue(topic.fields.filter(name='Nome do produto', parent__isnull=True).exists())
        budget_1 = topic.fields.get(name='Orçamento 1', parent__isnull=True)
        self.assertTrue(topic.fields.filter(parent=budget_1, name='Preço').exists())
        self.assertTrue(topic.fields.filter(parent=budget_1, name='Frete').exists())

    def test_can_create_topic(self):
        response = self.client.post(
            reverse('create_topic'),
            {'name': 'Material permanente'},
        )

        topic = CostTopic.objects.get(name='Material permanente')
        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}",
        )

    def test_can_create_field_and_subfield(self):
        topic = CostTopic.objects.create(name='Transporte')

        response = self.client.post(
            reverse('create_topic_field'),
            {
                'topic_id': topic.id,
                'name': 'Destino',
                'field_type': 'texto',
                'parent_id': '',
            },
        )
        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}&open=campos",
        )

        parent = CostField.objects.get(topic=topic, name='Destino')
        response = self.client.post(
            reverse('create_topic_field'),
            {
                'topic_id': topic.id,
                'name': 'Link da passagem',
                'field_type': 'link',
                'parent_id': parent.id,
            },
        )
        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}&open=campos",
        )

        child = CostField.objects.get(topic=topic, name='Link da passagem')
        self.assertEqual(child.parent, parent)

    def test_can_create_field_with_calculation_role(self):
        topic = CostTopic.objects.create(name='Material permanente')

        response = self.client.post(
            reverse('create_topic_field'),
            {
                'topic_id': topic.id,
                'name': 'Preço unitário',
                'field_type': 'valor',
                'calculation_role': CostField.ROLE_UNIT_PRICE,
                'parent_id': '',
            },
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}&open=campos",
        )
        field = CostField.objects.get(topic=topic, name='Preço unitário')
        self.assertEqual(field.calculation_role, CostField.ROLE_UNIT_PRICE)

    def test_field_form_includes_boolean_type(self):
        form = TopicFieldForm()

        self.assertIn(
            (CostField.TYPE_BOOLEAN, 'Sim/Não'),
            list(form.fields['field_type'].choices),
        )

    def test_campos_modal_starts_open_when_querystring_requests_it(self):
        topic = CostTopic.objects.create(name='Material permanente')

        response = self.client.get(
            reverse('budget_product_create'),
            {'topic': topic.id, 'open': 'campos'},
        )

        self.assertContains(response, 'class="modal is-open" id="campos-modal"', html=False)
        self.assertContains(response, 'aria-hidden="false"', html=False)

    def test_can_create_and_display_boolean_field_value(self):
        topic = CostTopic.objects.create(name='Critérios')
        field = CostField.objects.create(
            topic=topic,
            name='Aprovado',
            field_type=CostField.TYPE_BOOLEAN,
        )

        response = self.client.post(
            reverse('create_topic_record'),
            {
                'topic_id': topic.id,
                f'field_{field.id}': 'nao',
            },
        )

        self.assertRedirects(response, f"{reverse('budget_product_create')}?topic={topic.id}")
        saved_value = CostRecordValue.objects.get(field=field).value
        self.assertEqual(saved_value, 'nao')

        response = self.client.get(reverse('budget_product_create'), {'topic': topic.id})
        self.assertContains(response, 'Aprovado')
        self.assertContains(response, 'Não')

    def test_can_update_field_with_calculation_role(self):
        topic = CostTopic.objects.create(name='Material permanente')
        parent = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        field = CostField.objects.create(
            topic=topic,
            name='Campo antigo',
            field_type='texto',
        )

        response = self.client.post(
            reverse('update_topic_field', args=[field.id]),
            {
                'name': 'Preço ajustado',
                'field_type': 'valor',
                'calculation_role': CostField.ROLE_UNIT_PRICE,
                'parent_id': str(parent.id),
            },
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}&open=campos",
        )
        field.refresh_from_db()
        self.assertEqual(field.name, 'Preço ajustado')
        self.assertEqual(field.field_type, 'valor')
        self.assertEqual(field.calculation_role, CostField.ROLE_UNIT_PRICE)
        self.assertEqual(field.parent, parent)

    def test_can_create_dynamic_cost_record(self):
        topic = CostTopic.objects.create(name='Material permanente')
        product = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        quote = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='valor')
        link = CostField.objects.create(
            topic=topic,
            parent=quote,
            name='Link do orçamento',
            field_type='link',
        )

        response = self.client.post(
            reverse('create_topic_record'),
            {
                'topic_id': topic.id,
                f'field_{product.id}': 'Notebook Dell',
                f'field_{quote.id}': '4.500,90',
                f'field_{link.id}': 'https://example.com/notebook',
            },
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}",
        )
        record = CostRecord.objects.get(topic=topic)
        self.assertEqual(record.values.count(), 3)
        self.assertTrue(
            CostRecordValue.objects.filter(
                record=record,
                field=product,
                value='Notebook Dell',
            ).exists()
        )
        self.assertTrue(
            CostRecordValue.objects.filter(
                record=record,
                field=quote,
                value='4.500,90',
            ).exists()
        )
        self.assertTrue(
            CostRecordValue.objects.filter(
                record=record,
                field=link,
                value='https://example.com/notebook',
            ).exists()
        )

    def test_create_dynamic_cost_record_strips_utm_tracking_from_links(self):
        topic = CostTopic.objects.create(name='Material permanente')
        link_field = CostField.objects.create(topic=topic, name='Link do orçamento', field_type='link')

        response = self.client.post(
            reverse('create_topic_record'),
            {
                'topic_id': topic.id,
                f'field_{link_field.id}': 'https://example.com/produto?utm_source=chatgpt.com&utm_medium=referral&id=10',
            },
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}",
        )
        record = CostRecord.objects.get(topic=topic)
        saved_value = CostRecordValue.objects.get(record=record, field=link_field)
        self.assertEqual(saved_value.value, 'https://example.com/produto?id=10')

    def test_budget_page_shows_dynamic_builder(self):
        topic = CostTopic.objects.create(name='Bolsas')
        CostField.objects.create(topic=topic, name='Modalidade', field_type='texto')

        response = self.client.get(
            reverse('budget_product_create'),
            {'topic': topic.id},
        )

        self.assertContains(response, 'Monte os tópicos e registre os custos do orçamento')
        self.assertContains(response, 'Bolsas')
        self.assertContains(response, 'Modalidade')
        self.assertNotContains(response, 'tÃ³picos', html=False)
        self.assertNotContains(response, 'orÃ§amento', html=False)

    def test_budget_page_shows_search_and_alphabetic_order_controls_for_topic_records(self):
        topic = CostTopic.objects.create(name='Bolsas')
        field = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=field, value='Alpha')

        response = self.client.get(
            reverse('budget_product_create'),
            {'topic': topic.id},
        )

        self.assertContains(response, 'data-record-search="selected-topic-records"', html=False)
        self.assertContains(response, 'Comece a digitar o nome do cadastro')
        self.assertContains(response, 'A-Z')
        self.assertContains(response, 'data-record-list="selected-topic-records"', html=False)

    def test_budget_ready_reflects_dynamic_topics_and_records(self):
        topic = CostTopic.objects.create(
            name='Serviços de Terceiros contratados no país e no exterior',
            description='Resumo dinâmico do tópico.',
        )
        service_name = CostField.objects.create(
            topic=topic,
            name='Nome do serviço',
            field_type='texto',
        )
        selector = CostField.objects.create(
            topic=topic,
            name='Selecionar para orçar',
            field_type='numero',
        )
        quote_1 = CostField.objects.create(
            topic=topic,
            name='Orçamento 1',
            field_type='texto',
        )
        quote_price = CostField.objects.create(
            topic=topic,
            parent=quote_1,
            name='Preço',
            field_type='valor',
        )
        quote_freight = CostField.objects.create(
            topic=topic,
            parent=quote_1,
            name='Frete',
            field_type='valor',
        )
        quote_link = CostField.objects.create(
            topic=topic,
            parent=quote_1,
            name='Link',
            field_type='link',
        )

        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=service_name, value='Transcrição de entrevistas')
        CostRecordValue.objects.create(record=record, field=selector, value='1')
        CostRecordValue.objects.create(record=record, field=quote_price, value='120.00')
        CostRecordValue.objects.create(record=record, field=quote_freight, value='30.00')
        CostRecordValue.objects.create(record=record, field=quote_link, value='https://example.com/orcamento')

        response = self.client.get(reverse('budget_ready'))

        self.assertContains(response, 'Serviços de Terceiros contratados no país e no exterior')
        self.assertContains(response, 'Transcrição de entrevistas')
        self.assertContains(response, 'Resumo dinâmico do tópico.')
        self.assertContains(response, 'R$ 150,00', html=False)
        self.assertContains(response, 'Abrir link')
        self.assertContains(response, 'https://example.com/orcamento')
        self.assertContains(response, '120,00')

    def test_budget_ready_shows_search_and_alphabetic_order_controls_for_each_topic(self):
        topic = CostTopic.objects.create(name='Material permanente')
        field = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=field, value='Alpha')

        response = self.client.get(reverse('budget_ready'))

        self.assertContains(response, f'data-record-search="budget-topic-{topic.id}"', html=False)
        self.assertContains(response, 'A-Z')
        self.assertContains(response, f'data-record-list="budget-topic-{topic.id}"', html=False)

    def test_budget_ready_pdf_exports_file(self):
        topic = CostTopic.objects.create(name='Material permanente')
        field = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=field, value='Notebook')

        response = self.client.get(reverse('budget_ready_pdf'))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertIn('attachment; filename="orcamento-neevy.pdf"', response['Content-Disposition'])
        self.assertTrue(len(response.content) > 0)

    def test_budget_ready_selected_pdf_exports_file(self):
        topic = CostTopic.objects.create(name='Material permanente')
        selector = CostField.objects.create(
            topic=topic,
            name='Selecionar para orçar',
            field_type='numero',
            calculation_role=CostField.ROLE_SELECTOR,
        )
        budget = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        price = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Preço',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        quantity = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Quantidade',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        product = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=product, value='Notebook')
        CostRecordValue.objects.create(record=record, field=selector, value='1')
        CostRecordValue.objects.create(record=record, field=price, value='12,50')
        CostRecordValue.objects.create(record=record, field=quantity, value='2')

        response = self.client.get(reverse('budget_ready_selected_pdf'))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertIn(
            'attachment; filename="orcamentos-selecionados-neevy.pdf"',
            response['Content-Disposition'],
        )
        self.assertTrue(len(response.content) > 0)

    def test_build_fapesp_export_pages_maps_material_permanente(self):
        topic = CostTopic.objects.create(name='Material permanente adquirido no país e importado')
        selector = CostField.objects.create(
            topic=topic,
            name='Selecionar para orçar',
            field_type='numero',
            calculation_role=CostField.ROLE_SELECTOR,
        )
        budget = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        price = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Preço',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        quantity = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Quantidade',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        link = CostField.objects.create(topic=topic, parent=budget, name='Link', field_type='link')
        total = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Valor total',
            field_type='valor',
            calculation_role=CostField.ROLE_CALCULATED_TOTAL,
        )
        product = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=product, value='Notebook Lenovo')
        CostRecordValue.objects.create(record=record, field=selector, value='1')
        CostRecordValue.objects.create(record=record, field=price, value='1500,00')
        CostRecordValue.objects.create(record=record, field=quantity, value='2')
        CostRecordValue.objects.create(record=record, field=link, value='https://example.com/notebook')
        CostRecordValue.objects.create(record=record, field=total, value='3000,00')

        context = build_budget_ready_context()
        pages = build_fapesp_export_pages(context['topic_blocks'])
        page = next(item for item in pages if item['record_title'] == 'Notebook Lenovo')

        self.assertEqual(page['form_title'], 'Material Permanente')
        self.assertIn(('Origem *', 'Brasil'), page['rows'])
        self.assertIn(('Quantidade *', '2'), page['rows'])
        self.assertIn(('Fabricado no Brasil *', 'Sim'), page['rows'])
        self.assertIn(('Valor Unitário *', '1.500,00'), page['rows'])
        self.assertIn(('Valor Total *', '3.000,00'), page['rows'])
        description_row = next(value for label, value in page['rows'] if label == 'Descrição *')
        self.assertIn('Notebook Lenovo', description_row)
        self.assertIn('https://example.com/notebook', description_row)

    def test_build_fapesp_export_pages_infers_service_quantity_as_one(self):
        topic = CostTopic.objects.create(name='Serviços de Terceiros contratados no país e no exterior')
        selector = CostField.objects.create(
            topic=topic,
            name='Selecionar para orçar',
            field_type='numero',
            calculation_role=CostField.ROLE_SELECTOR,
        )
        budget = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        price = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Preço',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        service = CostField.objects.create(topic=topic, name='Nome do serviço', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=service, value='Assistência técnica')
        CostRecordValue.objects.create(record=record, field=selector, value='1')
        CostRecordValue.objects.create(record=record, field=price, value='2250,00')

        context = build_budget_ready_context()
        pages = build_fapesp_export_pages(context['topic_blocks'])
        page = next(item for item in pages if item['record_title'] == 'Assistência técnica')

        self.assertEqual(page['form_title'], 'Serviços de Terceiros')
        self.assertIn(('Quantidade *', '1'), page['rows'])
        self.assertIn(('Valor Unitário *', '2.250,00'), page['rows'])
        self.assertIn(('Valor Total *', '2.250,00'), page['rows'])

    def test_build_fapesp_export_pages_splits_event_from_transport(self):
        topic = CostTopic.objects.create(name='Despesas de Transporte e Diárias')
        selector = CostField.objects.create(
            topic=topic,
            name='Selecionar para orçar',
            field_type='numero',
            calculation_role=CostField.ROLE_SELECTOR,
        )
        budget = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        price = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Preço',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        quantity = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Quantidade de pesquisadores',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        event_name = CostField.objects.create(topic=topic, name='Nome do meio de transporte/Diárias/Eventos', field_type='texto')
        flight_name = CostField.objects.create(topic=topic, name='Origem - Destino', field_type='texto')

        event_record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=event_record, field=event_name, value='Evento internacional')
        CostRecordValue.objects.create(record=event_record, field=selector, value='1')
        CostRecordValue.objects.create(record=event_record, field=price, value='46500,00')
        CostRecordValue.objects.create(record=event_record, field=quantity, value='1')

        flight_record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=flight_record, field=flight_name, value='Voo de Fortaleza para Campinas')
        CostRecordValue.objects.create(record=flight_record, field=selector, value='1')
        CostRecordValue.objects.create(record=flight_record, field=price, value='1359,62')
        CostRecordValue.objects.create(record=flight_record, field=quantity, value='2')

        context = build_budget_ready_context()
        pages = build_fapesp_export_pages(context['topic_blocks'])
        event_page = next(item for item in pages if item['record_title'] == 'Evento internacional')
        flight_page = next(item for item in pages if item['record_title'] == 'Voo de Fortaleza para Campinas')

        self.assertEqual(event_page['form_title'], 'Diárias')
        self.assertIn(('Pernoite *', 'Sim'), event_page['rows'])
        self.assertEqual(flight_page['form_title'], 'Despesas de Transporte')

    def test_budget_ready_fapesp_pdf_exports_file(self):
        topic = CostTopic.objects.create(name='Material permanente')
        selector = CostField.objects.create(
            topic=topic,
            name='Selecionar para orçar',
            field_type='numero',
            calculation_role=CostField.ROLE_SELECTOR,
        )
        budget = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        price = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Preço',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        quantity = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Quantidade',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        product = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=product, value='Notebook')
        CostRecordValue.objects.create(record=record, field=selector, value='1')
        CostRecordValue.objects.create(record=record, field=price, value='12,50')
        CostRecordValue.objects.create(record=record, field=quantity, value='2')

        response = self.client.get(reverse('budget_ready_fapesp_pdf'))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/pdf')
        self.assertIn(
            'attachment; filename="modelo-fapesp-neevy.pdf"',
            response['Content-Disposition'],
        )
        self.assertTrue(len(response.content) > 0)

    def test_budget_ready_fapesp_docx_exports_file(self):
        topic = CostTopic.objects.create(name='Material permanente')
        selector = CostField.objects.create(
            topic=topic,
            name='Selecionar para orçar',
            field_type='numero',
            calculation_role=CostField.ROLE_SELECTOR,
        )
        budget = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        price = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Preço',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        quantity = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Quantidade',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        product = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=product, value='Notebook')
        CostRecordValue.objects.create(record=record, field=selector, value='1')
        CostRecordValue.objects.create(record=record, field=price, value='12,50')
        CostRecordValue.objects.create(record=record, field=quantity, value='2')

        response = self.client.get(reverse('budget_ready_fapesp_docx'))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        self.assertIn(
            'attachment; filename="modelo-fapesp-neevy.docx"',
            response['Content-Disposition'],
        )
        self.assertTrue(len(response.content) > 0)
        document = Document(BytesIO(response.content))
        paragraph_text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
        table_text = '\n'.join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        document_text = f'{paragraph_text}\n{table_text}'
        self.assertIn('Material Permanente', document_text)
        self.assertIn('Notebook', document_text)
        self.assertIn('Valor Unitário *', document_text)

    def test_budget_ready_docx_exports_file(self):
        topic = CostTopic.objects.create(name='Material permanente')
        field = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        link_group = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        link_field = CostField.objects.create(topic=topic, parent=link_group, name='Link', field_type='link')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=field, value='Notebook')
        CostRecordValue.objects.create(record=record, field=link_field, value='https://example.com/orcamento-1')

        response = self.client.get(reverse('budget_ready_docx'))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        self.assertIn('attachment; filename="orcamento-neevy.docx"', response['Content-Disposition'])
        self.assertTrue(len(response.content) > 0)
        document = Document(BytesIO(response.content))
        paragraph_text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
        table_text = '\n'.join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        document_text = f'{paragraph_text}\n{table_text}'
        self.assertIn('https://example.com/orcamento-1', document_text)

    def test_budget_ready_selected_docx_exports_file(self):
        topic = CostTopic.objects.create(name='Material permanente')
        selector = CostField.objects.create(
            topic=topic,
            name='Selecionar para or\u00e7ar',
            field_type='numero',
            calculation_role=CostField.ROLE_SELECTOR,
        )
        budget = CostField.objects.create(topic=topic, name='Or\u00e7amento 1', field_type='texto')
        price = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Pre\u00e7o',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        quantity = CostField.objects.create(
            topic=topic,
            parent=budget,
            name='Quantidade',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        product = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=product, value='Notebook')
        CostRecordValue.objects.create(record=record, field=selector, value='1')
        CostRecordValue.objects.create(record=record, field=price, value='12,50')
        CostRecordValue.objects.create(record=record, field=quantity, value='2')

        response = self.client.get(reverse('budget_ready_selected_docx'))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        self.assertIn(
            'attachment; filename="orcamentos-selecionados-neevy.docx"',
            response['Content-Disposition'],
        )
        self.assertTrue(len(response.content) > 0)
        document = Document(BytesIO(response.content))
        paragraph_text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
        table_text = '\n'.join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        document_text = f'{paragraph_text}\n{table_text}'
        self.assertIn('Notebook', document_text)
        self.assertIn('12,50', document_text)
        self.assertIn('25,00', document_text)

    def test_store_field_uses_suggestion_list_with_defaults_and_saved_values(self):
        topic = CostTopic.objects.create(name='Material permanente')
        store_field = CostField.objects.create(topic=topic, name='Loja', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=store_field, value='Kalunga')

        response = self.client.get(reverse('budget_product_create'), {'topic': topic.id})

        self.assertContains(response, 'list="store-suggestions-list"', html=False)
        self.assertContains(response, '<option value="Amazon">', html=False)
        self.assertContains(response, '<option value="Kabum">', html=False)
        self.assertContains(response, '<option value="Magazine Luiza">', html=False)
        self.assertContains(response, '<option value="Mercado Livre">', html=False)
        self.assertContains(response, '<option value="Shopee">', html=False)
        self.assertContains(response, '<option value="Kalunga">', html=False)

    def test_budget_totals_use_ptbr_thousands_separator(self):
        topic = CostTopic.objects.create(name='Material permanente')
        selector = CostField.objects.create(
            topic=topic,
            name='Selecionar para orçar',
            field_type='numero',
            calculation_role=CostField.ROLE_SELECTOR,
        )
        quote_1 = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        quote_price = CostField.objects.create(
            topic=topic,
            parent=quote_1,
            name='Preço',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        quote_quantity = CostField.objects.create(
            topic=topic,
            parent=quote_1,
            name='Quantidade',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )

        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=selector, value='1')
        CostRecordValue.objects.create(record=record, field=quote_price, value='70.306,27')
        CostRecordValue.objects.create(record=record, field=quote_quantity, value='1000')

        response = self.client.get(reverse('budget_product_create'), {'topic': topic.id})

        self.assertContains(response, 'R$&nbsp;70.306.270,00', html=True)
        self.assertContains(response, f'name="field_{selector.id}" value="1"', html=False)
        self.assertContains(response, f'name="field_{quote_quantity.id}" value="1"', html=False)
        self.assertContains(response, f'data-budget-selector="topic-{topic.id}"', html=False)
        self.assertContains(response, f'data-budget-total-group="topic-{topic.id}"', html=False)

    def test_budget_selector_ignores_zero_value_budgets_when_choosing_the_smallest(self):
        topic = CostTopic.objects.create(name='Material permanente')
        selector = CostField.objects.create(
            topic=topic,
            name='Selecionar para orÃ§ar',
            field_type='numero',
            calculation_role=CostField.ROLE_SELECTOR,
        )
        quote_1 = CostField.objects.create(topic=topic, name='OrÃ§amento 1', field_type='texto')
        CostField.objects.create(
            topic=topic,
            parent=quote_1,
            name='PreÃ§o',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        quote_2 = CostField.objects.create(topic=topic, name='OrÃ§amento 2', field_type='texto')
        CostField.objects.create(
            topic=topic,
            parent=quote_2,
            name='PreÃ§o',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )

        response = self.client.get(reverse('budget_product_create'), {'topic': topic.id})

        self.assertContains(response, f'data-budget-selector="topic-{topic.id}"', html=False)
        self.assertContains(response, 'totalValue <= 0', html=False)
        self.assertContains(response, f'name="field_{selector.id}" value="1"', html=False)

    def test_calculated_total_field_is_saved_from_price_and_multiplier(self):
        topic = CostTopic.objects.create(name='Material permanente')
        selector = CostField.objects.create(
            topic=topic,
            name='Selecionar para orçar',
            field_type='numero',
            calculation_role=CostField.ROLE_SELECTOR,
        )
        quote_1 = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        quote_price = CostField.objects.create(
            topic=topic,
            parent=quote_1,
            name='Preço',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        quote_quantity = CostField.objects.create(
            topic=topic,
            parent=quote_1,
            name='Quantidade',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        quote_freight = CostField.objects.create(
            topic=topic,
            parent=quote_1,
            name='Frete',
            field_type='valor',
            calculation_role=CostField.ROLE_FREIGHT,
        )
        quote_total = CostField.objects.create(
            topic=topic,
            parent=quote_1,
            name='Preço total',
            field_type='valor',
            calculation_role=CostField.ROLE_CALCULATED_TOTAL,
        )

        response = self.client.post(
            reverse('create_topic_record'),
            {
                'topic_id': topic.id,
                f'field_{selector.id}': '1',
                f'field_{quote_price.id}': '10,00',
                f'field_{quote_quantity.id}': '5',
                f'field_{quote_freight.id}': '2,00',
                f'field_{quote_total.id}': '',
            },
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}",
        )
        record = CostRecord.objects.get(topic=topic)
        total_value = CostRecordValue.objects.get(record=record, field=quote_total)
        self.assertEqual(total_value.value, '52,00')

    def test_standalone_calculation_supports_value_times_duration_times_quantity(self):
        topic = CostTopic.objects.create(name='Bolsas')
        value_field = CostField.objects.create(
            topic=topic,
            name='Valor orçamentário (por estudante)',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        quantity_field = CostField.objects.create(
            topic=topic,
            name='Quantidade de estudantes',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        duration_field = CostField.objects.create(
            topic=topic,
            name='Duração em meses',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        total_field = CostField.objects.create(
            topic=topic,
            name='Preço total',
            field_type='valor',
            calculation_role=CostField.ROLE_CALCULATED_TOTAL,
        )

        response = self.client.post(
            reverse('create_topic_record'),
            {
                'topic_id': topic.id,
                f'field_{value_field.id}': '500,00',
                f'field_{quantity_field.id}': '3',
                f'field_{duration_field.id}': '12',
                f'field_{total_field.id}': '',
            },
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}",
        )
        record = CostRecord.objects.get(topic=topic)
        total_value = CostRecordValue.objects.get(record=record, field=total_field)
        self.assertEqual(total_value.value, '18.000,00')

        response = self.client.get(reverse('budget_product_create'), {'topic': topic.id})
        self.assertContains(response, 'R$&nbsp;18.000,00', html=True)

    def test_can_edit_dynamic_cost_record_and_recalculate_total(self):
        topic = CostTopic.objects.create(name='Material permanente')
        product = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        price = CostField.objects.create(
            topic=topic,
            name='Preço',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        quantity = CostField.objects.create(
            topic=topic,
            name='Quantidade',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        total = CostField.objects.create(
            topic=topic,
            name='Preço total',
            field_type='valor',
            calculation_role=CostField.ROLE_CALCULATED_TOTAL,
        )

        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=product, value='Notebook')
        CostRecordValue.objects.create(record=record, field=price, value='10,00')
        CostRecordValue.objects.create(record=record, field=quantity, value='2')
        CostRecordValue.objects.create(record=record, field=total, value='20,00')

        response = self.client.post(
            reverse('update_topic_record', args=[record.id]),
            {
                f'field_{product.id}': 'Notebook atualizado',
                f'field_{price.id}': '15,00',
                f'field_{quantity.id}': '3',
                f'field_{total.id}': '',
            },
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}",
        )
        self.assertTrue(
            CostRecordValue.objects.filter(
                record=record,
                field=product,
                value='Notebook atualizado',
            ).exists()
        )
        total_value = CostRecordValue.objects.get(record=record, field=total)
        self.assertEqual(total_value.value, '45,00')
        self.assertTrue(
            AuditLog.objects.filter(
                action=AuditLog.ACTION_UPDATE,
                target_type='Custo',
                target_name='Notebook atualizado',
            ).exists()
        )

    def test_edit_dynamic_cost_record_strips_utm_tracking_from_links(self):
        topic = CostTopic.objects.create(name='Material permanente')
        link_field = CostField.objects.create(topic=topic, name='Link', field_type='link')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=link_field, value='https://example.com/original')

        response = self.client.post(
            reverse('update_topic_record', args=[record.id]),
            {
                f'field_{link_field.id}': 'https://example.com/novo?utm_source=chatgpt.com&sku=22',
            },
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}",
        )
        updated_value = CostRecordValue.objects.get(record=record, field=link_field)
        self.assertEqual(updated_value.value, 'https://example.com/novo?sku=22')

    def test_can_migrate_old_fields_to_new_field_in_single_record(self):
        topic = CostTopic.objects.create(name='Material permanente')
        old_field_1 = CostField.objects.create(topic=topic, name='Campo legado A', field_type='texto')
        old_field_2 = CostField.objects.create(topic=topic, name='Campo legado B', field_type='texto')
        new_field = CostField.objects.create(topic=topic, name='Resumo novo', field_type='texto')

        record_1 = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record_1, field=old_field_1, value='Primeira linha')
        CostRecordValue.objects.create(record=record_1, field=old_field_2, value='Segunda linha')

        record_2 = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record_2, field=old_field_1, value='Outro valor')
        CostRecordValue.objects.create(record=record_2, field=old_field_2, value='Mais um valor')

        response = self.client.post(
            reverse('migrate_topic_fields', args=[topic.id]),
            {
                'source_field_ids': [str(old_field_1.id), str(old_field_2.id)],
                'target_field_id': str(new_field.id),
                'record_id': str(record_1.id),
                'migration_mode': 'single',
            },
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}&open=registro-{record_1.id}",
        )
        migrated_value = CostRecordValue.objects.get(record=record_1, field=new_field)
        self.assertEqual(migrated_value.value, 'Primeira linha\nSegunda linha')
        self.assertFalse(CostRecordValue.objects.filter(record=record_2, field=new_field).exists())
        old_field_1.refresh_from_db()
        old_field_2.refresh_from_db()
        self.assertTrue(old_field_1.is_active)
        self.assertTrue(old_field_2.is_active)

    def test_field_migration_form_groups_fields_by_main_block(self):
        topic = CostTopic.objects.create(name='Material permanente')
        budget_root = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        CostField.objects.create(topic=topic, parent=budget_root, name='Preço', field_type='valor')
        CostField.objects.create(topic=topic, parent=budget_root, name='Quantidade', field_type='numero')
        standalone = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=standalone, value='Notebook')

        form = FieldMigrationForm(topic=topic, record_choices=[(str(record.id), 'Notebook')])
        rendered_source = str(form.fields['source_field_ids'].choices)

        self.assertIn('Orçamento 1', rendered_source)
        self.assertIn('Orçamento 1 / Preço', rendered_source)
        self.assertIn('Orçamento 1 / Quantidade', rendered_source)
        self.assertIn('Nome do produto', rendered_source)
        self.assertNotIn('Campo principal', rendered_source)

    def test_field_migration_target_shows_only_new_fields(self):
        topic = CostTopic.objects.create(name='Material permanente')
        old_root = CostField.objects.create(topic=topic, name='Orçamento 1', field_type='texto')
        old_field = CostField.objects.create(topic=topic, parent=old_root, name='Preço', field_type='valor')
        new_field = CostField.objects.create(topic=topic, name='Campo novo vazio', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=old_field, value='Valor antigo')

        form = FieldMigrationForm(topic=topic, record_choices=[(str(record.id), 'Teste')])

        target_choices = form.fields['target_field_id'].choices
        rendered_target = str(target_choices)
        rendered_source = str(form.fields['source_field_ids'].choices)

        self.assertIn('Orçamento 1 / Preço', rendered_source)
        self.assertIn('Campo novo vazio', rendered_source)
        self.assertNotIn('Orçamento 1', rendered_target)
        self.assertNotIn('Orçamento 1 / Preço', rendered_target)
        self.assertIn('Campo novo vazio', rendered_target)

    def test_can_apply_field_migration_to_all_records_and_archive_old_fields(self):
        topic = CostTopic.objects.create(name='Material permanente')
        old_field_1 = CostField.objects.create(topic=topic, name='Campo legado X1', field_type='texto')
        old_field_2 = CostField.objects.create(topic=topic, name='Campo legado X2', field_type='texto')
        new_field = CostField.objects.create(topic=topic, name='Resumo consolidado X', field_type='texto')

        record_1 = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record_1, field=old_field_1, value='Linha A')
        CostRecordValue.objects.create(record=record_1, field=old_field_2, value='Linha B')

        record_2 = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record_2, field=old_field_1, value='Linha C')
        CostRecordValue.objects.create(record=record_2, field=old_field_2, value='Linha D')

        response = self.client.post(
            reverse('migrate_topic_fields', args=[topic.id]),
            {
                'source_field_ids': [str(old_field_1.id), str(old_field_2.id)],
                'target_field_id': str(new_field.id),
                'archive_source_fields': 'on',
                'migration_mode': 'all',
            },
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}&open=campos",
        )
        self.assertEqual(
            CostRecordValue.objects.get(record=record_1, field=new_field).value,
            'Linha A\nLinha B',
        )
        self.assertEqual(
            CostRecordValue.objects.get(record=record_2, field=new_field).value,
            'Linha C\nLinha D',
        )
        old_field_1.refresh_from_db()
        old_field_2.refresh_from_db()
        self.assertFalse(old_field_1.is_active)
        self.assertFalse(old_field_2.is_active)

        response = self.client.get(reverse('budget_ready'))
        self.assertContains(response, 'Resumo consolidado X')
        self.assertContains(response, 'Linha A')
        self.assertNotContains(response, 'Campo legado X1', html=False)
        self.assertNotContains(response, 'Campo legado X2', html=False)

    def test_can_transfer_dynamic_cost_record_between_compatible_topics(self):
        source_topic = CostTopic.objects.create(name='Material permanente')
        source_product = CostField.objects.create(topic=source_topic, name='Nome do produto', field_type='texto')
        source_budget = CostField.objects.create(topic=source_topic, name='Orçamento 1', field_type='texto')
        source_price = CostField.objects.create(
            topic=source_topic,
            parent=source_budget,
            name='Preço',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        source_quantity = CostField.objects.create(
            topic=source_topic,
            parent=source_budget,
            name='Quantidade',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        source_total = CostField.objects.create(
            topic=source_topic,
            parent=source_budget,
            name='Preço total',
            field_type='valor',
            calculation_role=CostField.ROLE_CALCULATED_TOTAL,
        )

        target_topic = CostTopic.objects.create(name='Material de consumo')
        target_product = CostField.objects.create(topic=target_topic, name='Nome do produto', field_type='texto')
        target_budget = CostField.objects.create(topic=target_topic, name='Orçamento 1', field_type='texto')
        target_price = CostField.objects.create(
            topic=target_topic,
            parent=target_budget,
            name='Preço',
            field_type='valor',
            calculation_role=CostField.ROLE_UNIT_PRICE,
        )
        target_quantity = CostField.objects.create(
            topic=target_topic,
            parent=target_budget,
            name='Quantidade',
            field_type='numero',
            calculation_role=CostField.ROLE_MULTIPLIER,
        )
        target_total = CostField.objects.create(
            topic=target_topic,
            parent=target_budget,
            name='Preço total',
            field_type='valor',
            calculation_role=CostField.ROLE_CALCULATED_TOTAL,
        )

        record = CostRecord.objects.create(topic=source_topic)
        CostRecordValue.objects.create(record=record, field=source_product, value='Notebook Dell')
        CostRecordValue.objects.create(record=record, field=source_price, value='10,00')
        CostRecordValue.objects.create(record=record, field=source_quantity, value='2')
        CostRecordValue.objects.create(record=record, field=source_total, value='20,00')

        response = self.client.post(
            reverse('transfer_topic_record', args=[record.id]),
            {'target_topic_id': str(target_topic.id)},
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={target_topic.id}&open=registro-{record.id}",
        )
        record.refresh_from_db()
        self.assertEqual(record.topic, target_topic)
        self.assertTrue(
            CostRecordValue.objects.filter(
                record=record,
                field=target_product,
                value='Notebook Dell',
            ).exists()
        )
        self.assertTrue(
            CostRecordValue.objects.filter(
                record=record,
                field=target_price,
                value='10,00',
            ).exists()
        )
        self.assertTrue(
            CostRecordValue.objects.filter(
                record=record,
                field=target_quantity,
                value='2',
            ).exists()
        )
        self.assertTrue(
            CostRecordValue.objects.filter(
                record=record,
                field=target_total,
                value='20,00',
            ).exists()
        )
        self.assertEqual(record.values.filter(field__topic=source_topic).count(), 0)
        self.assertTrue(
            AuditLog.objects.filter(
                action=AuditLog.ACTION_UPDATE,
                target_type='Custo',
                description__contains='Transferência de custo do tópico Material permanente para Material de consumo.',
            ).exists()
        )

    def test_transfer_dynamic_cost_record_is_blocked_without_compatible_fields(self):
        source_topic = CostTopic.objects.create(name='Material permanente')
        source_product = CostField.objects.create(topic=source_topic, name='Nome do produto', field_type='texto')
        record = CostRecord.objects.create(topic=source_topic)
        CostRecordValue.objects.create(record=record, field=source_product, value='Notebook Dell')

        target_topic = CostTopic.objects.create(name='Transporte')
        CostField.objects.create(topic=target_topic, name='Destino', field_type='texto')

        response = self.client.post(
            reverse('transfer_topic_record', args=[record.id]),
            {'target_topic_id': str(target_topic.id)},
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={source_topic.id}&open=transferir-registro-{record.id}",
        )
        record.refresh_from_db()
        self.assertEqual(record.topic, source_topic)
        self.assertTrue(
            CostRecordValue.objects.filter(
                record=record,
                field=source_product,
                value='Notebook Dell',
            ).exists()
        )

    def test_can_delete_dynamic_cost_record(self):
        topic = CostTopic.objects.create(name='Material permanente')
        field = CostField.objects.create(topic=topic, name='Nome do produto', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=field, value='Projetor')

        response = self.client.post(
            reverse('delete_topic_record', args=[record.id]),
            {'next': f"{reverse('budget_product_create')}?topic={topic.id}"},
        )

        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}",
        )
        self.assertFalse(CostRecord.objects.filter(id=record.id).exists())

    def test_can_delete_topic_with_records(self):
        topic = CostTopic.objects.create(name='Transporte')
        field = CostField.objects.create(topic=topic, name='Destino', field_type='texto')
        record = CostRecord.objects.create(topic=topic)
        CostRecordValue.objects.create(record=record, field=field, value='São Paulo')

        response = self.client.post(reverse('delete_topic', args=[topic.id]))

        self.assertRedirects(response, reverse('budget_product_create'))
        self.assertFalse(CostTopic.objects.filter(id=topic.id).exists())
        self.assertFalse(CostRecord.objects.filter(id=record.id).exists())


class AuditLogTests(TestCase):
    def test_admin_dashboard_shows_audit_button(self):
        self.client.post(reverse('login'), {'username': 'adm', 'password': '123'})

        response = self.client.get(reverse('dashboard'))

        self.assertContains(response, 'Auditoria')

    def test_non_admin_cannot_access_audit_page(self):
        self.client.post(reverse('login'), {'username': 'fabiano', 'password': '123'})

        response = self.client.get(reverse('audit_log'))

        self.assertRedirects(response, reverse('dashboard'))

    def test_creating_topic_generates_audit_log(self):
        self.client.post(reverse('login'), {'username': 'adm', 'password': '123'})

        response = self.client.post(reverse('create_topic'), {'name': 'Tópico auditado'})

        topic = CostTopic.objects.get(name='Tópico auditado')
        self.assertRedirects(
            response,
            f"{reverse('budget_product_create')}?topic={topic.id}",
        )
        audit_log = AuditLog.objects.get(
            action=AuditLog.ACTION_CREATE,
            target_type='Tópico',
            target_name='Tópico auditado',
        )
        self.assertEqual(audit_log.user.login_name, 'adm')

    def test_audit_page_lists_registered_actions(self):
        admin_user = User.objects.get(login_name='adm')
        AuditLog.objects.create(
            user=admin_user,
            action=AuditLog.ACTION_DELETE,
            target_type='Custo',
            target_name='Notebook',
            description='Exclusão de custo no tópico Material permanente.',
        )
        self.client.post(reverse('login'), {'username': 'adm', 'password': '123'})

        response = self.client.get(reverse('audit_log'))

        self.assertContains(response, 'Notebook')
        self.assertContains(response, 'Exclusão')
        self.assertContains(response, 'adm')


class AllowedSignupEmailTests(TestCase):
    def test_admin_can_view_allowed_signup_emails_page(self):
        self.client.post(reverse('login'), {'username': 'adm', 'password': '123'})

        response = self.client.get(reverse('allowed_signup_emails'))

        self.assertContains(response, 'E-mails autorizados para cadastro')

    def test_admin_can_add_dynamic_signup_email(self):
        self.client.post(reverse('login'), {'username': 'adm', 'password': '123'})

        response = self.client.post(
            reverse('create_allowed_signup_email'),
            {'email': 'novo.email@example.com'},
        )

        self.assertRedirects(response, reverse('allowed_signup_emails'))
        self.assertTrue(AllowedSignupEmail.objects.filter(email='novo.email@example.com').exists())
        self.assertTrue(
            AuditLog.objects.filter(
                action=AuditLog.ACTION_CREATE,
                target_type='E-mail autorizado',
                target_name='novo.email@example.com',
            ).exists()
        )

    def test_admin_can_remove_dynamic_signup_email(self):
        self.client.post(reverse('login'), {'username': 'adm', 'password': '123'})
        allowed_email = AllowedSignupEmail.objects.create(email='remover@example.com')

        response = self.client.post(reverse('delete_allowed_signup_email', args=[allowed_email.id]))

        self.assertRedirects(response, reverse('allowed_signup_emails'))
        self.assertFalse(AllowedSignupEmail.objects.filter(email='remover@example.com').exists())

    def test_dynamic_signup_email_is_accepted_in_signup_flow(self):
        AllowedSignupEmail.objects.create(email='liberado@example.com')

        response = self.client.post(reverse('signup_email'), {'email': 'liberado@example.com'})

        self.assertRedirects(response, reverse('login'))
        self.assertEqual(SignupCode.objects.filter(email='liberado@example.com').count(), 1)
