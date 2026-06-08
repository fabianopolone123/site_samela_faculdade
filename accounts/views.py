import secrets
from decimal import Decimal, InvalidOperation
from io import BytesIO
from urllib.parse import parse_qsl, urlencode, urlsplit, urlunsplit
from xml.sax.saxutils import escape

from django.conf import settings
from django.contrib import messages
from django.contrib.auth import logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.views import LoginView
from django.core.mail import EmailMessage
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .forms import (
    AllowedSignupEmailForm,
    LoginForm,
    SignupCodeForm,
    SignupEmailForm,
    SignupPasswordForm,
    TopicDescriptionForm,
    TopicFieldForm,
    TopicForm,
)
from .models import (
    AuditLog,
    AllowedSignupEmail,
    CostField,
    CostRecord,
    CostRecordValue,
    CostTopic,
    SignupCode,
    User,
)

SIGNUP_EMAIL_SESSION_KEY = 'signup_email'
SIGNUP_STEP_SESSION_KEY = 'signup_step'
SIGNUP_CODE_SESSION_KEY = 'signup_code_id'
DEFAULT_STORE_SUGGESTIONS = [
    'Mercado Livre',
    'Kabum',
    'Shopee',
    'Amazon',
    'Magazine Luiza',
]
PROJECT_INFO = {
    'title': 'Orçamento - FAPESP - Fundação Bracell Fundação Itaú',
    'subtitle': 'Auxílio à Pesquisa para o Fortalecimento da Educação na Pré-Escola',
    'edition': '06/2026',
    'organization': 'NEEVY - UFSCar',
}
PROJECT_BUDGET_TITLE = (
    'INDICADORES E CRITÉRIOS DE AVALIAÇÃO DE DESENVOLVIMENTO CULTURAL '
    'DE CRIANÇAS DE PRÉ-ESCOLA NA THC'
)
PROJECT_BUDGET_DESCRIPTION = (
    'A FAPESP, conforme estabelecido no convênio celebrado com a Fundação Bracell '
    'e a Fundação Itaú, cobrirá os custos do projeto de pesquisa segundo normas e '
    'orientações para Auxílio à Pesquisa Regular (para projetos com teto orçamentário '
    'de R$ 600 mil) ou para Projeto Temático (sem teto orçamentário). Serão aprovadas '
    'propostas até o limite orçamentário da Chamada (total de R$ 6.400.000,00). '
    'O orçamento do projeto de pesquisa apresentado à FAPESP deverá ser detalhado e '
    'cada item justificado especificamente em termos dos objetivos do projeto proposto.'
)


def _send_utf8_mail(subject, body, to):
    msg = EmailMessage(
        subject=subject,
        body=body,
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=to,
    )
    msg.encoding = 'utf-8'
    msg.send(fail_silently=False)


class ProjectLoginView(LoginView):
    authentication_form = LoginForm
    template_name = 'accounts/login.html'
    redirect_authenticated_user = True

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context.update(build_signup_context(self.request))
        context['project_info'] = PROJECT_INFO
        return context


def login_view(request):
    return ProjectLoginView.as_view()(request)


@login_required
def dashboard_view(request):
    return render(request, 'accounts/dashboard.html')


@login_required
def allowed_signup_emails_view(request):
    if not request.user.is_staff:
        messages.error(request, 'A gestão de e-mails está disponível apenas para o login de administração.')
        return redirect('dashboard')

    dynamic_emails = list(AllowedSignupEmail.objects.all())
    dynamic_values = {item.email for item in dynamic_emails}
    fixed_emails = sorted(settings.ALLOWED_SIGNUP_EMAILS - dynamic_values)
    context = {
        'email_form': AllowedSignupEmailForm(),
        'fixed_emails': fixed_emails,
        'dynamic_emails': dynamic_emails,
    }
    return render(request, 'accounts/allowed_signup_emails.html', context)


@login_required
def create_allowed_signup_email_view(request):
    if request.method != 'POST':
        return redirect('allowed_signup_emails')
    if not request.user.is_staff:
        messages.error(request, 'A gestão de e-mails está disponível apenas para o login de administração.')
        return redirect('dashboard')

    form = AllowedSignupEmailForm(request.POST)
    if not form.is_valid():
        messages.error(request, 'Informe um e-mail válido para liberar o cadastro.')
        return redirect('allowed_signup_emails')

    email = normalize_email(form.cleaned_data['email'])
    if email in settings.ALLOWED_SIGNUP_EMAILS or AllowedSignupEmail.objects.filter(email=email).exists():
        messages.info(request, 'Este e-mail já está autorizado para cadastro.')
        return redirect('allowed_signup_emails')

    AllowedSignupEmail.objects.create(email=email)
    register_audit_log(
        request.user,
        AuditLog.ACTION_CREATE,
        'E-mail autorizado',
        email,
        'Liberação de novo e-mail para o fluxo de cadastro.',
    )
    messages.success(request, 'E-mail autorizado cadastrado com sucesso.')
    return redirect('allowed_signup_emails')


@login_required
def delete_allowed_signup_email_view(request, email_id):
    if request.method != 'POST':
        return redirect('allowed_signup_emails')
    if not request.user.is_staff:
        messages.error(request, 'A gestão de e-mails está disponível apenas para o login de administração.')
        return redirect('dashboard')

    allowed_email = get_object_or_404(AllowedSignupEmail, id=email_id)
    email = allowed_email.email
    allowed_email.delete()
    register_audit_log(
        request.user,
        AuditLog.ACTION_DELETE,
        'E-mail autorizado',
        email,
        'Remoção de e-mail da lista dinâmica de cadastro.',
    )
    messages.success(request, 'E-mail autorizado removido com sucesso.')
    return redirect('allowed_signup_emails')


@login_required
def audit_log_view(request):
    if not request.user.is_staff:
        messages.error(request, 'A auditoria está disponível apenas para o login de administração.')
        return redirect('dashboard')

    audit_logs = AuditLog.objects.select_related('user').all()[:300]
    return render(request, 'accounts/audit_log.html', {'audit_logs': audit_logs})


def build_budget_ready_context():
    all_topics = list(
        CostTopic.objects.prefetch_related('fields', 'records__values__field').all()
    )
    topic_blocks = []
    all_topics_total = Decimal('0')

    for topic in all_topics:
        rows = build_topic_rows(topic)
        records, topic_total = build_record_cards(topic, rows)
        topic_blocks.append(
            {
                'topic': topic,
                'records': records,
                'topic_total': topic_total,
            }
        )
        all_topics_total += topic_total

    return {
        'project_budget_title': PROJECT_BUDGET_TITLE,
        'project_budget_description': PROJECT_BUDGET_DESCRIPTION,
        'topic_blocks': topic_blocks,
        'all_topics_total': all_topics_total,
    }


@login_required
def budget_product_create_view(request):
    topics = list(
        CostTopic.objects.prefetch_related(
            'fields',
            'records__values__field',
        )
    )

    selected_topic = None
    selected_topic_id = request.GET.get('topic')
    if selected_topic_id:
        selected_topic = next(
            (topic for topic in topics if str(topic.id) == selected_topic_id),
            None,
        )
    elif topics:
        selected_topic = topics[0]

    selected_rows = build_topic_rows(selected_topic) if selected_topic else []
    field_edit_forms = build_field_edit_forms(selected_topic)
    for row in selected_rows:
        row['edit_form'] = field_edit_forms.get(row['field'].id)
    selected_groups = build_topic_groups(selected_topic) if selected_topic else []
    has_standalone_calculation = any(
        not group['children'] and group['root_role'] in (
            CostField.ROLE_UNIT_PRICE,
            CostField.ROLE_MULTIPLIER,
            CostField.ROLE_FREIGHT,
            CostField.ROLE_CALCULATED_TOTAL,
        )
        for group in selected_groups
    )
    selected_records, topic_grand_total = (
        build_record_cards(selected_topic, selected_rows) if selected_topic else ([], Decimal('0'))
    )

    all_topics_total = sum(calculate_topic_total(topic) for topic in topics)
    store_suggestions = get_store_suggestions()

    context = {
        'topics': topics,
        'selected_topic': selected_topic,
        'topic_form': TopicForm(),
        'field_form': TopicFieldForm(topic=selected_topic),
        'description_form': TopicDescriptionForm(
            initial={'description': selected_topic.description} if selected_topic else {}
        ),
        'selected_rows': selected_rows,
        'selected_groups': selected_groups,
        'has_standalone_calculation': has_standalone_calculation,
        'selected_records': selected_records,
        'topic_grand_total': topic_grand_total,
        'all_topics_total': all_topics_total,
        'store_suggestions': store_suggestions,
    }
    return render(request, 'accounts/budget_product_form.html', context)


@login_required
def create_topic_view(request):
    if request.method != 'POST':
        return redirect('budget_product_create')

    form = TopicForm(request.POST)
    if form.is_valid():
        topic = CostTopic.objects.create(name=form.cleaned_data['name'])
        register_audit_log(
            request.user,
            AuditLog.ACTION_CREATE,
            'Tópico',
            topic.name,
            'Cadastro de novo tópico na área de custos.',
        )
        messages.success(request, 'Tópico cadastrado com sucesso.')
        return redirect(f"{reverse('budget_product_create')}?topic={topic.id}")

    messages.error(request, 'Não foi possível cadastrar o tópico.')
    return redirect('budget_product_create')


@login_required
def delete_topic_view(request, topic_id):
    if request.method != 'POST':
        return redirect('budget_product_create')

    topic = get_object_or_404(CostTopic, id=topic_id)
    topic_name = topic.name
    topic.delete()
    register_audit_log(
        request.user,
        AuditLog.ACTION_DELETE,
        'Tópico',
        topic_name,
        'Exclusão do tópico e dos registros vinculados.',
    )
    messages.success(request, 'Tópico excluído com sucesso.')
    return redirect('budget_product_create')


@login_required
def create_topic_field_view(request):
    if request.method != 'POST':
        return redirect('budget_product_create')

    topic = get_object_or_404(CostTopic, id=request.POST.get('topic_id'))
    form = TopicFieldForm(request.POST, topic=topic)
    if form.is_valid():
        parent = None
        parent_id = form.cleaned_data.get('parent_id')
        if parent_id:
            parent = get_object_or_404(CostField, id=parent_id, topic=topic)

        CostField.objects.create(
            topic=topic,
            parent=parent,
            name=form.cleaned_data['name'],
            field_type=form.cleaned_data['field_type'],
            calculation_role=form.cleaned_data['calculation_role'] or CostField.ROLE_NONE,
        )
        register_audit_log(
            request.user,
            AuditLog.ACTION_CREATE,
            'Campo',
            form.cleaned_data['name'],
            f'Cadastro de campo no tópico {topic.name}.',
        )
        messages.success(request, 'Campo cadastrado com sucesso.')
    else:
        messages.error(request, 'Não foi possível cadastrar o campo.')

    return redirect(f"{reverse('budget_product_create')}?topic={topic.id}&open=campos")


@login_required
def update_topic_field_view(request, field_id):
    if request.method != 'POST':
        return redirect('budget_product_create')

    field = get_object_or_404(CostField, id=field_id)
    topic = field.topic
    form = TopicFieldForm(request.POST, topic=topic, current_field=field)
    if form.is_valid():
        parent = None
        parent_id = form.cleaned_data.get('parent_id')
        if parent_id:
            parent = get_object_or_404(CostField, id=parent_id, topic=topic)

        previous_name = field.name
        field.name = form.cleaned_data['name']
        field.field_type = form.cleaned_data['field_type']
        field.calculation_role = form.cleaned_data['calculation_role'] or CostField.ROLE_NONE
        field.parent = parent
        field.save()
        register_audit_log(
            request.user,
            AuditLog.ACTION_UPDATE,
            'Campo',
            field.name,
            f'Edição do campo {previous_name} no tópico {topic.name}.',
        )
        messages.success(request, 'Campo atualizado com sucesso.')
    else:
        messages.error(request, 'Não foi possível atualizar o campo.')

    return redirect(f"{reverse('budget_product_create')}?topic={topic.id}&open=campos")


@login_required
def update_topic_description_view(request, topic_id):
    if request.method != 'POST':
        return redirect('budget_product_create')
    topic = get_object_or_404(CostTopic, id=topic_id)
    form = TopicDescriptionForm(request.POST)
    if form.is_valid():
        topic.description = form.cleaned_data['description']
        topic.save()
        register_audit_log(
            request.user,
            AuditLog.ACTION_UPDATE,
            'Tópico',
            topic.name,
            'Atualização da observação do tópico.',
        )
        messages.success(request, 'Observação salva com sucesso.')
    else:
        messages.error(request, 'Não foi possível salvar a observação.')
    return redirect(f"{reverse('budget_product_create')}?topic={topic_id}")


@login_required
def delete_topic_field_view(request, field_id):
    if request.method != 'POST':
        return redirect('budget_product_create')

    field = get_object_or_404(CostField, id=field_id)
    topic_id = field.topic_id
    field_name = field.name
    topic_name = field.topic.name
    field.delete()
    register_audit_log(
        request.user,
        AuditLog.ACTION_DELETE,
        'Campo',
        field_name,
        f'Exclusão de campo do tópico {topic_name}.',
    )
    messages.success(request, 'Campo excluído com sucesso.')
    return redirect(f"{reverse('budget_product_create')}?topic={topic_id}&open=campos")


@login_required
def create_topic_record_view(request):
    if request.method != 'POST':
        return redirect('budget_product_create')

    topic = get_object_or_404(CostTopic, id=request.POST.get('topic_id'))
    fields = list(topic.fields.order_by('created_at'))

    if not fields:
        messages.error(request, 'Crie ao menos um campo antes de salvar um novo custo.')
        return redirect(f"{reverse('budget_product_create')}?topic={topic.id}")

    record = CostRecord.objects.create(topic=topic)
    saved_values = 0
    raw_values = {}

    for field in fields:
        value = request.POST.get(f'field_{field.id}', '').strip()
        if field.field_type == 'valor' and value:
            parsed_value = parse_decimal_input(value)
            if parsed_value is not None:
                value = format_decimal_br(parsed_value)
        elif field.field_type == 'link' and value:
            value = sanitize_external_url(value)
        raw_values[field.id] = value

    calculated_totals = {}
    for field in fields:
        if get_effective_calculation_role(field) != CostField.ROLE_CALCULATED_TOTAL:
            continue
        calculated_totals[field.id] = format_decimal_br(
            get_group_calculation_parts(fields, raw_values, field.parent_id)['total']
        )

    for field in fields:
        value = calculated_totals.get(field.id, raw_values.get(field.id, ''))
        if value:
            CostRecordValue.objects.create(record=record, field=field, value=value)
            saved_values += 1

    if saved_values == 0:
        record.delete()
        messages.error(request, 'Preencha ao menos um campo para salvar o novo custo.')
    else:
        register_audit_log(
            request.user,
            AuditLog.ACTION_CREATE,
            'Custo',
            get_record_title_from_record(record),
            f'Cadastro de novo custo no tópico {topic.name}.',
        )
        messages.success(request, 'Novo custo cadastrado com sucesso.')

    return redirect(f"{reverse('budget_product_create')}?topic={topic.id}")


@login_required
def update_topic_record_view(request, record_id):
    if request.method != 'POST':
        return redirect('budget_product_create')

    record = get_object_or_404(CostRecord.objects.select_related('topic'), id=record_id)
    topic = record.topic
    fields = list(topic.fields.order_by('created_at'))

    raw_values = {}
    for field in fields:
        value = request.POST.get(f'field_{field.id}', '').strip()
        if field.field_type == 'valor' and value:
            parsed_value = parse_decimal_input(value)
            if parsed_value is not None:
                value = format_decimal_br(parsed_value)
        elif field.field_type == 'link' and value:
            value = sanitize_external_url(value)
        raw_values[field.id] = value

    calculated_totals = {}
    for field in fields:
        if get_effective_calculation_role(field) != CostField.ROLE_CALCULATED_TOTAL:
            continue
        calculated_totals[field.id] = format_decimal_br(
            get_group_calculation_parts(fields, raw_values, field.parent_id)['total']
        )

    values_to_save = {}
    for field in fields:
        value = calculated_totals.get(field.id, raw_values.get(field.id, ''))
        if value:
            values_to_save[field.id] = value

    if not values_to_save:
        messages.error(request, 'Preencha ao menos um campo para salvar as alterações do custo.')
        return redirect(f"{reverse('budget_product_create')}?topic={topic.id}&open=registro-{record.id}")

    record.values.all().delete()
    CostRecordValue.objects.bulk_create(
        [
            CostRecordValue(record=record, field_id=field_id, value=value)
            for field_id, value in values_to_save.items()
        ]
    )
    register_audit_log(
        request.user,
        AuditLog.ACTION_UPDATE,
        'Custo',
        get_record_title_from_record(record),
        f'Edição de custo no tópico {topic.name}.',
    )
    messages.success(request, 'Custo atualizado com sucesso.')
    return redirect(f"{reverse('budget_product_create')}?topic={topic.id}")


@login_required
def delete_topic_record_view(request, record_id):
    if request.method != 'POST':
        return redirect('budget_product_create')

    record = get_object_or_404(CostRecord, id=record_id)
    topic_id = record.topic_id
    topic_name = record.topic.name
    record_title = get_record_title_from_record(record)
    next_url = request.POST.get('next', '').strip()
    record.delete()
    register_audit_log(
        request.user,
        AuditLog.ACTION_DELETE,
        'Custo',
        record_title,
        f'Exclusão de custo do tópico {topic_name}.',
    )
    messages.success(request, 'Custo excluído com sucesso.')

    if next_url.startswith('/'):
        return redirect(next_url)
    return redirect(f"{reverse('budget_product_create')}?topic={topic_id}")


@login_required
def budget_ready_view(request):
    context = build_budget_ready_context()
    return render(request, 'accounts/budget_ready.html', context)


@login_required
def budget_ready_pdf_view(request):
    context = build_budget_ready_context()
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title='Orçamento NEEVY',
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'BudgetPdfTitle',
        parent=styles['Title'],
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#0b5563'),
        spaceAfter=8,
    )
    section_title_style = ParagraphStyle(
        'BudgetPdfSectionTitle',
        parent=styles['Heading2'],
        fontSize=12,
        leading=15,
        textColor=colors.HexColor('#172126'),
        spaceBefore=8,
        spaceAfter=4,
    )
    small_style = ParagraphStyle(
        'BudgetPdfSmall',
        parent=styles['BodyText'],
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#62727f'),
    )
    body_style = ParagraphStyle(
        'BudgetPdfBody',
        parent=styles['BodyText'],
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#172126'),
        wordWrap='CJK',
    )
    value_style = ParagraphStyle(
        'BudgetPdfValue',
        parent=styles['BodyText'],
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#172126'),
        alignment=2,
    )
    link_style = ParagraphStyle(
        'BudgetPdfLink',
        parent=styles['BodyText'],
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#0b5563'),
        wordWrap='CJK',
    )

    story = [
        Paragraph('Orçamento pronto', small_style),
        Paragraph(f'PROJETO: {context["project_budget_title"]}', title_style),
        Paragraph('ORÇAMENTO', section_title_style),
        Paragraph(context['project_budget_description'], body_style),
        Spacer(1, 8),
        Paragraph(
            f'Total geral do projeto: R$ {format_decimal_br(context["all_topics_total"])}',
            section_title_style,
        ),
        Spacer(1, 8),
    ]

    for block in context['topic_blocks']:
        story.append(Paragraph(block['topic'].name, section_title_style))
        if block['topic'].description:
            story.append(Paragraph(block['topic'].description, small_style))
        story.append(
            Paragraph(
                f'Total do tópico: R$ {format_decimal_br(block["topic_total"])}',
                body_style,
            )
        )
        story.append(Spacer(1, 6))

        if not block['records']:
            story.append(Paragraph('Nenhum item cadastrado neste tópico.', small_style))
            story.append(Spacer(1, 8))
            continue

        for record_card in block['records']:
            story.append(Paragraph(record_card['record_title'], body_style))
            story.append(
                Paragraph(
                    f'Cadastrado em {record_card["record"].created_at.strftime("%d/%m/%Y %H:%M")}',
                    small_style,
                )
            )
            if record_card['selected_total'] is not None:
                story.append(
                    Paragraph(
                        f'Total do orçamento selecionado: R$ {format_decimal_br(record_card["selected_total"])}',
                        small_style,
                    )
                )

            summary_rows = []
            for detail_group in record_card['detail_groups']:
                if detail_group['is_simple_field']:
                    if detail_group['root_is_url']:
                        value = build_pdf_link_value(detail_group['root_raw_value'], link_style, 'Abrir link')
                    else:
                        value = build_pdf_text_value(detail_group['summary_value'], body_style)
                    summary_rows.append([
                        build_pdf_text_value(detail_group['title'], body_style),
                        value or build_pdf_text_value('-', body_style),
                    ])

            if summary_rows:
                table = Table(summary_rows, colWidths=[55 * mm, 115 * mm])
                table.setStyle(
                    TableStyle(
                        [
                            ('BACKGROUND', (0, 0), (-1, -1), colors.whitesmoke),
                            ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#d8e3e4')),
                            ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#d8e3e4')),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('LEFTPADDING', (0, 0), (-1, -1), 6),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                            ('TOPPADDING', (0, 0), (-1, -1), 5),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 6))

            for detail_group in record_card['detail_groups']:
                if not detail_group['children']:
                    continue
                budget_heading = detail_group['title']
                if detail_group['is_selected_budget']:
                    budget_heading += ' (selecionado)'
                story.append(Paragraph(budget_heading, small_style))
                budget_rows = []
                for item in detail_group['children']:
                    if item['is_url_value']:
                        value = build_pdf_link_value(item['raw_value'], link_style, 'Abrir orçamento')
                    else:
                        value = build_pdf_text_value(item['display_value'], body_style)
                    budget_rows.append([
                        build_pdf_text_value(item['field_name'], body_style),
                        value or build_pdf_text_value('-', body_style),
                    ])
                if detail_group['group_total'] is not None:
                    budget_rows.append(['Total deste orçamento', f'R$ {format_decimal_br(detail_group["group_total"])}'])
                budget_table = Table(budget_rows, colWidths=[55 * mm, 115 * mm])
                budget_table.setStyle(
                    TableStyle(
                        [
                            ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#d8e3e4')),
                            ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#d8e3e4')),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('LEFTPADDING', (0, 0), (-1, -1), 6),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                            ('TOPPADDING', (0, 0), (-1, -1), 5),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                        ]
                    )
                )
                story.append(budget_table)
                story.append(Spacer(1, 6))
            story.append(Spacer(1, 8))

    doc.build(story)
    pdf = buffer.getvalue()
    buffer.close()

    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="orcamento-neevy.pdf"'
    return response


@login_required
def budget_ready_docx_view(request):
    context = build_budget_ready_context()
    document = Document()

    document.core_properties.title = 'Orçamento NEEVY'
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('Orçamento pronto')
    run.bold = True
    run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run(f'PROJETO: {context["project_budget_title"]}')
    run.bold = True
    run.font.size = Pt(12)

    document.add_paragraph('ORÇAMENTO').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph(context['project_budget_description'])
    total_paragraph = document.add_paragraph()
    total_run = total_paragraph.add_run(
        f'Total geral do projeto: R$ {format_decimal_br(context["all_topics_total"])}'
    )
    total_run.bold = True

    for block in context['topic_blocks']:
        document.add_heading(block['topic'].name, level=2)
        if block['topic'].description:
            document.add_paragraph(block['topic'].description)
        topic_total = document.add_paragraph()
        topic_total.add_run(
            f'Total do tópico: R$ {format_decimal_br(block["topic_total"])}'
        ).bold = True

        if not block['records']:
            document.add_paragraph('Nenhum item cadastrado neste tópico.')
            continue

        for record_card in block['records']:
            record_heading = document.add_paragraph()
            record_heading.add_run(record_card['record_title']).bold = True
            document.add_paragraph(
                f'Cadastrado em {record_card["record"].created_at.strftime("%d/%m/%Y %H:%M")}'
            )
            if record_card['selected_total'] is not None:
                document.add_paragraph(
                    f'Total do orçamento selecionado: R$ {format_decimal_br(record_card["selected_total"])}'
                )

            summary_groups = [group for group in record_card['detail_groups'] if group['is_simple_field']]
            if summary_groups:
                document.add_paragraph('Dados principais').runs[0].bold = True
                summary_table = document.add_table(rows=0, cols=2)
                summary_table.style = 'Table Grid'
                for detail_group in summary_groups:
                    cells = summary_table.add_row().cells
                    cells[0].text = detail_group['title']
                    if detail_group['root_is_url']:
                        cells[1].text = detail_group['root_raw_value'] or '-'
                    else:
                        cells[1].text = detail_group['summary_value'] or '-'

            budget_groups = [group for group in record_card['detail_groups'] if group['children']]
            for detail_group in budget_groups:
                budget_paragraph = document.add_paragraph()
                budget_paragraph.add_run(detail_group['title']).bold = True
                if detail_group['is_selected_budget']:
                    budget_paragraph.add_run(' (selecionado)')
                budget_table = document.add_table(rows=0, cols=2)
                budget_table.style = 'Table Grid'
                for item in detail_group['children']:
                    cells = budget_table.add_row().cells
                    cells[0].text = item['field_name']
                    if item['is_url_value']:
                        paragraph = cells[1].paragraphs[0]
                        cells[1].text = item['raw_value'] or '-'
                        cells[1].text = item['raw_value'] or '-'
                        add_docx_hyperlink(paragraph, item['raw_value'], 'Abrir orçamento')
                    else:
                        cells[1].text = item['display_value'] or '-'
                if detail_group['group_total'] is not None:
                    cells = budget_table.add_row().cells
                    cells[0].text = 'Total deste orçamento'
                    cells[1].text = f'R$ {format_decimal_br(detail_group["group_total"])}'
            document.add_paragraph()

    buffer = BytesIO()
    document.save(buffer)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = 'attachment; filename="orcamento-neevy.docx"'
    return response


def logout_view(request):
    logout(request)
    messages.success(request, 'Sessão encerrada com sucesso.')
    return redirect('login')


def signup_email_view(request):
    if request.method != 'POST':
        return redirect('login')

    form = SignupEmailForm(request.POST)
    if not form.is_valid():
        persist_signup_state(request, step='email')
        return render_login_with_forms(request, email_form=form)

    email = normalize_email(form.cleaned_data['email'])

    if email not in get_allowed_signup_emails():
        form.add_error('email', 'Este e-mail não está autorizado para cadastro.')
        persist_signup_state(request, step='email')
        return render_login_with_forms(request, email_form=form)

    if User.objects.filter(email=email).exists():
        form.add_error('email', 'Já existe uma conta criada para este e-mail.')
        persist_signup_state(request, step='email')
        return render_login_with_forms(request, email_form=form)

    signup_code = SignupCode.objects.create(
        email=email,
        code=generate_code(),
        expires_at=SignupCode.expiration_time(),
    )

    _send_utf8_mail(
        subject='Seu código de acesso ao portal',
        body=(
            f'Seu código de verificação é: {signup_code.code}\n\n'
            f'Validade: {settings.SIGNUP_CODE_EXPIRATION_MINUTES} minutos.\n'
            'Se você não solicitou este cadastro, ignore esta mensagem.'
        ),
        to=[email],
    )

    persist_signup_state(
        request,
        email=email,
        step='code',
        code_id=signup_code.id,
    )
    messages.success(
        request,
        (
            f'Código enviado para {email}. '
            'Se não encontrar na caixa de entrada, verifique também Spam e Lixo eletrônico.'
        ),
    )
    return redirect('login')


def signup_code_view(request):
    if request.method != 'POST':
        return redirect('login')

    signup_code = get_active_signup_code(request)
    if signup_code is None:
        messages.error(request, 'Solicite um novo código para continuar.')
        clear_signup_state(request)
        return redirect('login')

    form = SignupCodeForm(request.POST)
    if not form.is_valid():
        persist_signup_state(
            request,
            email=signup_code.email,
            step='code',
            code_id=signup_code.id,
        )
        return render_login_with_forms(request, code_form=form)

    submitted_code = form.cleaned_data['code'].strip()
    if submitted_code != signup_code.code:
        form.add_error('code', 'Código inválido. Verifique o e-mail informado.')
        persist_signup_state(
            request,
            email=signup_code.email,
            step='code',
            code_id=signup_code.id,
        )
        return render_login_with_forms(request, code_form=form)

    signup_code.mark_verified()
    persist_signup_state(
        request,
        email=signup_code.email,
        step='password',
        code_id=signup_code.id,
    )
    messages.success(request, 'Código validado. Agora defina sua senha.')
    return redirect('login')


def signup_password_view(request):
    if request.method != 'POST':
        return redirect('login')

    signup_code = get_active_signup_code(request)
    if signup_code is None or signup_code.verified_at is None:
        messages.error(request, 'Valide o código antes de criar sua senha.')
        clear_signup_state(request)
        return redirect('login')

    form = SignupPasswordForm(request.POST)
    if not form.is_valid():
        persist_signup_state(
            request,
            email=signup_code.email,
            step='password',
            code_id=signup_code.id,
        )
        return render_login_with_forms(request, password_form=form)

    user = User.objects.create_user(
        email=signup_code.email,
        login_name=signup_code.email.split('@')[0],
        password=form.cleaned_data['password1'],
    )
    signup_code.mark_consumed()

    _send_utf8_mail(
        subject='Conta criada com sucesso no portal NEEVY - UFSCar',
        body=(
            'Seu cadastro foi concluído com sucesso.\n\n'
            f'E-mail de acesso: {user.email}\n'
            f'Data de criação: {timezone.localtime().strftime("%d/%m/%Y %H:%M")}\n'
            'A partir de agora você já pode fazer login no portal.'
        ),
        to=[user.email],
    )

    clear_signup_state(request)
    messages.success(request, 'Conta criada com sucesso. Faça login para continuar.')
    return redirect('login')


def signup_reset_view(request):
    clear_signup_state(request)
    messages.info(request, 'Fluxo de cadastro cancelado.')
    return redirect('login')


def render_login_with_forms(request, email_form=None, code_form=None, password_form=None):
    context = build_signup_context(request)
    context.update(
        {
            'form': LoginForm(request, data=request.POST if request.path == '/' else None),
            'email_form': email_form or SignupEmailForm(),
            'code_form': code_form or SignupCodeForm(),
            'password_form': password_form or SignupPasswordForm(),
            'project_info': PROJECT_INFO,
        }
    )
    return render(request, 'accounts/login.html', context)


def build_signup_context(request):
    email = request.session.get(SIGNUP_EMAIL_SESSION_KEY, '')
    step = request.session.get(SIGNUP_STEP_SESSION_KEY, 'email')
    return {
        'email_form': SignupEmailForm(initial={'email': email}),
        'code_form': SignupCodeForm(),
        'password_form': SignupPasswordForm(),
        'signup_email': email,
        'signup_step': step,
        'signup_modal_open': (
            SIGNUP_STEP_SESSION_KEY in request.session
            or SIGNUP_EMAIL_SESSION_KEY in request.session
        ),
    }


def persist_signup_state(request, email=None, step='email', code_id=None):
    if email:
        request.session[SIGNUP_EMAIL_SESSION_KEY] = email
    request.session[SIGNUP_STEP_SESSION_KEY] = step
    if code_id:
        request.session[SIGNUP_CODE_SESSION_KEY] = code_id
    request.session.modified = True


def clear_signup_state(request):
    for key in [
        SIGNUP_EMAIL_SESSION_KEY,
        SIGNUP_STEP_SESSION_KEY,
        SIGNUP_CODE_SESSION_KEY,
    ]:
        request.session.pop(key, None)
    request.session.modified = True


def get_active_signup_code(request):
    code_id = request.session.get(SIGNUP_CODE_SESSION_KEY)
    if not code_id:
        return None

    try:
        signup_code = SignupCode.objects.get(id=code_id)
    except SignupCode.DoesNotExist:
        return None

    if not signup_code.is_available:
        return None

    return signup_code


def build_topic_rows(topic):
    if topic is None:
        return []

    fields = list(topic.fields.all())
    children_map = {}
    for field in fields:
        children_map.setdefault(field.parent_id, []).append(field)

    rows = []

    def walk(field, level):
        rows.append(
            {
                'field': field,
                'level': level,
                'type_label': field.get_field_type_display(),
                'is_store_field': is_store_field(field),
            }
        )
        for child in children_map.get(field.id, []):
            walk(child, level + 1)

    for root in children_map.get(None, []):
        walk(root, 0)

    return rows


def build_field_edit_forms(topic):
    if topic is None:
        return {}

    forms_map = {}
    for field in topic.fields.all():
        forms_map[field.id] = TopicFieldForm(
            topic=topic,
            current_field=field,
            initial={
                'name': field.name,
                'field_type': field.field_type,
                'calculation_role': field.calculation_role,
                'parent_id': str(field.parent_id) if field.parent_id else '',
            },
        )
    return forms_map


def build_topic_groups(topic):
    if topic is None:
        return []

    rows = build_topic_rows(topic)
    groups = []
    current_group = None

    for row in rows:
        if row['level'] == 0:
            root_role = get_effective_calculation_role(row['field'])
            current_group = {
                'root': row,
                'root_role': root_role,
                'root_role_label': get_calculation_role_label(root_role),
                'children': [],
                'has_price_calc': False,
                'budget_number': extract_budget_number(row['field'].name),
            }
            groups.append(current_group)
            if root_role in (
                CostField.ROLE_UNIT_PRICE,
                CostField.ROLE_MULTIPLIER,
                CostField.ROLE_FREIGHT,
                CostField.ROLE_CALCULATED_TOTAL,
            ):
                current_group['has_price_calc'] = True
        elif current_group is not None:
            field_role = get_effective_calculation_role(row['field'])
            if field_role in (
                CostField.ROLE_UNIT_PRICE,
                CostField.ROLE_MULTIPLIER,
                CostField.ROLE_FREIGHT,
                CostField.ROLE_CALCULATED_TOTAL,
            ):
                current_group['has_price_calc'] = True
            current_group['children'].append(
                {
                    **row,
                    'field_role': field_role,
                    'field_role_label': get_calculation_role_label(field_role),
                }
            )

    return groups


def calculate_topic_total(topic):
    all_fields = list(topic.fields.all())
    total = Decimal('0')
    for record in topic.records.prefetch_related('values__field').all():
        topic_total = get_selected_total_for_record(record, all_fields)
        if topic_total is not None:
            total += topic_total
    return total


def get_selected_total_for_record(record, topic_fields):
    values_map = {record_value.field_id: record_value.value for record_value in record.values.all()}

    selector = next(
        (
            field for field in topic_fields
            if get_effective_calculation_role(field) == CostField.ROLE_SELECTOR
        ),
        None,
    )
    if not selector:
        calculation_parts = get_group_calculation_parts(topic_fields, values_map, None)
        if calculation_parts['has_parts']:
            return calculation_parts['total']
        return None

    selected_str = values_map.get(selector.id, '').strip()
    if not selected_str:
        return None

    budget_parent = next(
        (
            field for field in topic_fields
            if field.parent_id is None and field.name.strip() == f'Orçamento {selected_str}'
        ),
        None,
    )
    if not budget_parent:
        return None

    calculation_parts = get_group_calculation_parts(topic_fields, values_map, budget_parent.id)

    return calculation_parts['total']


def build_record_cards(topic, selected_rows):
    field_order = {row['field'].id: index for index, row in enumerate(selected_rows)}
    all_fields = list(topic.fields.all())
    selected_groups = build_topic_groups(topic)
    records = []
    grand_total = Decimal('0')

    for index, record in enumerate(topic.records.all(), start=1):
        values_by_field_id = {value.field_id: value.value for value in record.values.all()}
        values = sorted(
            [
                {
                    'field_name': value.field.name,
                    'field_type': value.field.get_field_type_display(),
                    'field_type_code': value.field.field_type,
                    'value': value.value,
                    'display_value': format_value_for_display(value.field.field_type, value.value),
                    'is_url_value': is_probably_url(value.value),
                    'level': get_field_level(value.field),
                    'order': field_order.get(value.field_id, 9999),
                }
                for value in record.values.all()
            ],
            key=lambda item: item['order'],
        )
        record_title = get_record_title(values, index)
        selected_total = get_selected_total_for_record(record, all_fields)
        if selected_total is not None:
            grand_total += selected_total
        records.append(
            {
                'record': record,
                'record_title': record_title,
                'values': values,
                'detail_groups': build_record_detail_groups(selected_groups, values_by_field_id, all_fields),
                'edit_groups': build_record_edit_groups(selected_groups, values_by_field_id),
                'selected_total': selected_total,
            }
        )

    return records, grand_total


def get_record_title(values, index):
    preferred_keywords = (
        'nome do produto',
        'nome do serviço',
        'nome do meio de transporte',
        'modalidade da bolsa',
        'modalidade',
        'origem',
    )
    for item in values:
        field_name = item['field_name'].lower()
        if any(keyword in field_name for keyword in preferred_keywords) and item['value'].strip():
            return item['value']

    for item in values:
        if item['value'].strip():
            return item['value']

    return f'Cadastro {index}'


def build_record_detail_groups(selected_groups, values_by_field_id, topic_fields):
    selected_budget_value = ''
    for group in selected_groups:
        root_field = group['root']['field']
        if group['root_role'] == CostField.ROLE_SELECTOR:
            selected_budget_value = str(values_by_field_id.get(root_field.id, '')).strip()
            break

    detail_groups = []
    for group in selected_groups:
        root_field = group['root']['field']
        root_value = values_by_field_id.get(root_field.id, '')
        child_items = []
        for child in group['children']:
            field_id = child['field'].id
            value = values_by_field_id.get(field_id, '')
            if value:
                child_items.append(
                    {
                        'field_name': child['field'].name,
                        'field_type_code': child['field'].field_type,
                        'display_value': format_value_for_display(child['field'].field_type, value),
                        'raw_value': value,
                        'is_url_value': is_probably_url(value),
                        'field_role': child['field_role'],
                    }
                )

        if root_value or child_items:
            budget_number = extract_budget_number(root_field.name)
            calculation_parts = get_group_calculation_parts(topic_fields, values_by_field_id, root_field.id)
            detail_groups.append(
                {
                    'title': root_field.name,
                    'type_label': group['root']['type_label'],
                    'root_value': format_value_for_display(root_field.field_type, root_value) if root_value else '',
                    'root_raw_value': root_value,
                    'root_is_url': is_probably_url(root_value) if root_value else False,
                    'root_type_code': root_field.field_type,
                    'children': child_items,
                    'is_budget_group': bool(group['children']) and budget_number is not None,
                    'is_selected_budget': bool(budget_number) and selected_budget_value == budget_number,
                    'budget_number': budget_number,
                    'group_total': calculation_parts['total'] if calculation_parts['has_parts'] else None,
                    'is_simple_field': not group['children'],
                    'summary_value': format_value_for_display(root_field.field_type, root_value) if root_value else '',
                }
            )
    return detail_groups


def build_record_edit_groups(selected_groups, values_by_field_id):
    edit_groups = []
    for group in selected_groups:
        root_field = group['root']['field']
        children = []
        for child in group['children']:
            children.append(
                {
                    **child,
                    'current_value': values_by_field_id.get(child['field'].id, ''),
                }
            )

        edit_groups.append(
            {
                **group,
                'root_current_value': values_by_field_id.get(root_field.id, ''),
                'children': children,
            }
        )
    return edit_groups


def get_record_title_from_record(record):
    values = [
        {
            'field_name': value.field.name,
            'value': value.value,
        }
        for value in record.values.select_related('field').all()
    ]
    return get_record_title(values, record.id or 0)


def extract_budget_number(field_name):
    normalized = field_name.lower().strip()
    if not normalized.startswith('orçamento ') and not normalized.startswith('orcamento '):
        return None

    parts = field_name.strip().split()
    if not parts:
        return None
    candidate = parts[-1].strip()
    return candidate if candidate.isdigit() else None


def get_field_level(field):
    level = 0
    current = field.parent
    while current is not None:
        level += 1
        current = current.parent
    return level


def is_store_field(field):
    normalized = field.name.lower().strip()
    return normalized == 'loja' or normalized.startswith('loja ')


def get_store_suggestions():
    suggestions = {name for name in DEFAULT_STORE_SUGGESTIONS}
    store_values = CostRecordValue.objects.select_related('field').all()
    for record_value in store_values:
        if not is_store_field(record_value.field):
            continue
        value = record_value.value.strip()
        if value:
            suggestions.add(value)
    return sorted(suggestions, key=lambda item: item.lower())


def generate_code():
    return f'{secrets.randbelow(1000000):06d}'


def normalize_email(email):
    return email.strip().lower()


def get_allowed_signup_emails():
    dynamic_emails = set(AllowedSignupEmail.objects.values_list('email', flat=True))
    return set(settings.ALLOWED_SIGNUP_EMAILS) | dynamic_emails


def get_effective_calculation_role(field):
    if field.calculation_role != CostField.ROLE_NONE:
        return field.calculation_role

    normalized = field.name.lower().strip()
    if 'selecionar' in normalized and 'orça' in normalized:
        return CostField.ROLE_SELECTOR
    if normalized in ('preço', 'preco', 'price'):
        return CostField.ROLE_UNIT_PRICE
    if normalized == 'quantidade':
        return CostField.ROLE_MULTIPLIER
    if normalized == 'frete':
        return CostField.ROLE_FREIGHT
    if normalized in ('total', 'preço total', 'preco total', 'valor total'):
        return CostField.ROLE_CALCULATED_TOTAL
    return CostField.ROLE_NONE


def get_calculation_role_label(role):
    role_map = dict(CostField.CALCULATION_ROLE_CHOICES)
    return role_map.get(role, role_map.get(CostField.ROLE_NONE, 'Sem função'))


def get_group_calculation_parts(topic_fields, values_map, parent_id):
    price_value = Decimal('0')
    multiplier_value = Decimal('1')
    freight_value = Decimal('0')
    has_multiplier = False
    has_parts = False

    for field in topic_fields:
        if field.parent_id != parent_id:
            continue
        value_str = str(values_map.get(field.id, '')).strip()
        if not value_str:
            continue
        amount = parse_decimal_input(value_str)
        if amount is None:
            continue

        role = get_effective_calculation_role(field)
        if role == CostField.ROLE_UNIT_PRICE:
            price_value = amount
            has_parts = True
        elif role == CostField.ROLE_MULTIPLIER:
            multiplier_value *= amount
            has_multiplier = True
            has_parts = True
        elif role == CostField.ROLE_FREIGHT:
            freight_value += amount
            has_parts = True
        elif role == CostField.ROLE_CALCULATED_TOTAL:
            has_parts = True

    if not has_multiplier:
        multiplier_value = Decimal('1')

    return {
        'price': price_value,
        'multiplier': multiplier_value,
        'freight': freight_value,
        'total': (price_value * multiplier_value) + freight_value,
        'has_parts': has_parts,
    }


def parse_decimal_input(value_str):
    cleaned = value_str.strip().replace('R$', '').replace(' ', '')
    if not cleaned:
        return None

    if ',' in cleaned:
        cleaned = cleaned.replace('.', '').replace(',', '.')

    try:
        return Decimal(cleaned)
    except (InvalidOperation, ValueError):
        return None


def format_decimal_br(value):
    rendered = f'{value:,.2f}'
    return rendered.replace(',', 'X').replace('.', ',').replace('X', '.')


def build_pdf_link_value(url, style, label):
    safe_url = escape(url, {'"': '&quot;'})
    safe_label = escape(label)
    return Paragraph(f'<link href="{safe_url}">{safe_label}</link>', style)


def build_pdf_text_value(value, style):
    return Paragraph(escape(str(value or '-')), style)


def add_docx_hyperlink(paragraph, url, text):
    return None


def sanitize_external_url(url):
    raw_url = (url or '').strip()
    if not raw_url:
        return raw_url

    try:
        parts = urlsplit(raw_url)
    except ValueError:
        return raw_url

    cleaned_query = [
        (key, value)
        for key, value in parse_qsl(parts.query, keep_blank_values=True)
        if not key.lower().startswith('utm_')
    ]
    return urlunsplit(
        (
            parts.scheme,
            parts.netloc,
            parts.path,
            urlencode(cleaned_query, doseq=True),
            parts.fragment,
        )
    )


def format_value_for_display(field_type_code, value):
    if field_type_code == 'valor':
        parsed = parse_decimal_input(value)
        if parsed is not None:
            return format_decimal_br(parsed)
    return value


def is_probably_url(value):
    normalized = value.strip().lower()
    return normalized.startswith('http://') or normalized.startswith('https://')


def register_audit_log(user, action, target_type, target_name, description=''):
    AuditLog.objects.create(
        user=user if getattr(user, 'is_authenticated', False) else None,
        action=action,
        target_type=target_type,
        target_name=target_name,
        description=description,
    )
